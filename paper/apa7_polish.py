# -*- coding: utf-8 -*-
"""Post-process the pandoc DOCX into an APA-7-styled, bilingual manuscript.

Applies: Times New Roman + SimSun 12pt, 1.5 line spacing, 2.54cm margins,
centered title/author/affiliation blocks, justified body, hanging-indent
reference list, bordered tables, centered figure + caption.
"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"F:\MCP\MNE-MCP\paper\_base.docx"
DST = r"F:\MCP\MNE-MCP\paper\MNE-MCP_论文_APA7.docx"

LATIN = "Times New Roman"
EASTASIA = "SimSun"  # 宋体
SIZE = Pt(12)

doc = Document(SRC)


def set_run_fonts(run):
    run.font.name = LATIN
    run.font.size = SIZE
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), LATIN)
    rfonts.set(qn("w:hAnsi"), LATIN)
    rfonts.set(qn("w:cs"), LATIN)
    rfonts.set(qn("w:eastAsia"), EASTASIA)


def style_eastasia(style):
    """Force east-asian + latin fonts on a style's rPr."""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), LATIN)
    rfonts.set(qn("w:hAnsi"), LATIN)
    rfonts.set(qn("w:cs"), LATIN)
    rfonts.set(qn("w:eastAsia"), EASTASIA)


# 1. Base style: Normal
normal = doc.styles["Normal"]
normal.font.name = LATIN
normal.font.size = SIZE
style_eastasia(normal)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(0)

# Also patch docDefaults so unstyled runs inherit the east-asian font.
styles_el = doc.styles.element
docdef = styles_el.find(qn("w:docDefaults"))
if docdef is not None:
    rpd = docdef.find(qn("w:rPrDefault"))
    if rpd is not None:
        rpr = rpd.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            rpd.append(rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), LATIN)
        rfonts.set(qn("w:hAnsi"), LATIN)
        rfonts.set(qn("w:eastAsia"), EASTASIA)

# 2. Margins 2.54 cm
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

paras = doc.paragraphs
n = len(paras)

# Locate the references heading.
ref_start = None
for i, p in enumerate(paras):
    if p.text.strip() == "参考文献":
        ref_start = i
        break

# Locate english title block start.
eng_title_idx = None
for i, p in enumerate(paras):
    if p.text.startswith("MNE-MCP: An Automated"):
        eng_title_idx = i
        break

center_idx = set()
# Chinese title block = first 3 non-empty paragraphs (title, authors, affiliation)
seen = 0
for i, p in enumerate(paras):
    if p.text.strip():
        center_idx.add(i)
        seen += 1
    if seen >= 3:
        break
# English title block = eng_title_idx and next 2 non-empty
if eng_title_idx is not None:
    cnt = 0
    for i in range(eng_title_idx, n):
        if paras[i].text.strip():
            center_idx.add(i)
            cnt += 1
        if cnt >= 3:
            break


def is_heading(text):
    t = text.strip()
    if not t:
        return False
    # Numbered section headers: "1 引言", "2.1 整体架构", "4 结论"
    head = t.split("　")[0].split(" ")[0]
    if head and head[0].isdigit():
        return True
    if t in ("参考文献",):
        return True
    return False


for i, p in enumerate(paras):
    txt = p.text.strip()
    # ensure runs carry fonts
    for r in p.runs:
        set_run_fonts(r)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    if i in center_idx:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        continue

    if ref_start is not None and i > ref_start and txt:
        # Reference entry -> hanging indent
        pf.left_indent = Cm(1.06)
        pf.first_line_indent = Cm(-1.06)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        continue

    if is_heading(txt):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(8)
        pf.space_after = Pt(4)
        continue

    # Figure caption / table caption -> center the figure-related captions
    if txt.startswith("图 1") or txt.startswith("图1"):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        continue

    # Body prose -> justify (longer paragraphs); leave short/quote lines as-is
    if len(txt) >= 30:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Center any paragraph that contains only the figure image (no text, has a drawing)
for p in paras:
    if not p.text.strip():
        if p._element.findall(".//" + qn("w:drawing")):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_table_borders(table):
    tbl = table._element
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")  # 0.5 pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


# 3. Tables: borders + fonts + spacing
for table in doc.tables:
    set_table_borders(table)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run_fonts(r)
                    r.font.size = Pt(10.5)

doc.save(DST)
print("WROTE", DST)
print("paragraphs:", n, "| ref_start:", ref_start, "| eng_title_idx:", eng_title_idx,
      "| tables:", len(doc.tables))
